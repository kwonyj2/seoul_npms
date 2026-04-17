"""
네트워크 서비스 – 토폴로지 빌드/저장
"""
from collections import defaultdict, deque

# ── 색상 테이블 ──────────────────────────────────────────────────
_NET_COLOR = {
    '교사망': '#0d6efd',
    '학생망': '#198754',
    '무선망': '#dc3545',
    '기타망': '#fd7e14',
    '전화망': '#6f42c1',
}
_TYPE_COLOR = {
    'firewall':   '#8b0000',
    'router':     '#fd7e14',
    'server':     '#6f42c1',
    'ap':         '#20c997',
    'poe_switch': '#0dcaf0',
    'switch':     '#0d6efd',
}
_CABLE_COLOR = {
    'fiber':   '#dc3545',
    'cat6':    '#0d6efd',
    'cat5e':   '#198754',
    'cat5':    '#6c757d',
    'unknown': '#adb5bd',
}
_CABLE_KO = {
    'fiber': '광', 'cat6': 'Cat6', 'cat5e': 'Cat5e', 'cat5': 'Cat5', 'unknown': '미확인'
}
_TYPE_LEVEL = {
    'firewall': 0, 'router': 1, 'server': 1,
    'switch': 2, 'poe_switch': 3, 'ap': 4,
}


def _compute_levels(devices, links):
    """링크 구조 BFS → 각 장비의 계층(level) 자동 계산"""
    all_ids  = {d.id for d in devices}
    children = defaultdict(list)
    has_parent = set()

    for lk in links:
        children[lk.from_device_id].append(lk.to_device_id)
        has_parent.add(lk.to_device_id)

    # 루트: 방화벽/라우터 우선, 없으면 incoming 없는 노드
    root_ids = {d.id for d in devices if d.device_type in ('firewall', 'router')}
    if not root_ids:
        root_ids = all_ids - has_parent
    if not root_ids:
        root_ids = all_ids

    levels = {}
    q = deque()
    for rid in root_ids:
        levels[rid] = 0
        q.append(rid)
    while q:
        nid = q.popleft()
        for cid in children[nid]:
            if cid not in levels:
                levels[cid] = levels[nid] + 1
                q.append(cid)

    # 링크 없는 고립 노드 → device_type 기반 fallback
    for d in devices:
        if d.id not in levels:
            levels[d.id] = _TYPE_LEVEL.get(d.device_type, 2)

    return levels


def build_topology_data(school_id: int) -> dict:
    """현재 DB 장비/링크 → vis.js 형식 토폴로지 반환"""
    from apps.network.models import NetworkDevice, NetworkLink

    devices = list(NetworkDevice.objects.filter(school_id=school_id))
    links   = list(NetworkLink.objects.filter(
        from_device__school_id=school_id, is_active=True
    ).select_related('from_device', 'to_device'))

    node_levels = _compute_levels(devices, links)

    # 장비 유형별 아이콘 경로
    ICON_MAP = {
        'firewall':   '/npms/static/img/network-icons/firewall.svg',
        'router':     '/npms/static/img/network-icons/router.svg',
        'l3_switch':  '/npms/static/img/network-icons/l3_switch.svg',
        'switch':     '/npms/static/img/network-icons/switch.svg',
        'l2_switch':  '/npms/static/img/network-icons/switch.svg',
        'poe_switch': '/npms/static/img/network-icons/poe_switch.svg',
        'ap':         '/npms/static/img/network-icons/ap.svg',
        'server':     '/npms/static/img/network-icons/server.svg',
    }

    nodes = []
    for d in devices:
        net    = d.network_type or ''
        color  = _NET_COLOR.get(net) or _TYPE_COLOR.get(d.device_type, '#0d6efd')
        is_root = d.device_type in ('firewall', 'router')
        level   = node_levels.get(d.id, _TYPE_LEVEL.get(d.device_type, 2))
        icon   = ICON_MAP.get(d.device_type, ICON_MAP['switch'])
        # 아이콘 아래 표시 라벨 (장비명 / 모델 / 위치)
        label_lines = [d.name]
        if d.model: label_lines.append(d.model)
        if d.location: label_lines.append(d.location)
        label = '\n'.join(label_lines)

        nodes.append({
            'id':          d.id,
            'label':       label,
            'title': (
                f'<b>{d.name}</b><br>'
                f'모델: {d.model or "-"}<br>'
                f'위치: {d.location or "-"}<br>'
                f'망: {net or "-"}<br>'
                f'IP: {d.ip_address or "미등록"}<br>'
                f'상태: {d.get_status_display()}'
            ),
            'shape': 'image',
            'image': icon,
            'size':  34 if is_root else 28,
            'font': {
                'size': 11,
                'color': '#263238',
                'face': 'Malgun Gothic, 맑은 고딕, sans-serif',
                'vadjust': 5,        # 아이콘과 라벨 간격
                'multi': 'html',
                'bold': {'color': color, 'size': 12, 'face': 'Malgun Gothic'},
            },
            'level': level,
            'network_type': net,
            'device_type':  d.device_type,
            'ip_address':   d.ip_address or '',
            'status':       d.status,
            'model':        d.model or '',
            'location':     d.location or '',
            'color':        {'border': color, 'background': color},
        })

    edges = []
    for lk in links:
        cable     = lk.cable_type or 'unknown'
        from_net  = lk.from_device.network_type or ''
        clr       = _NET_COLOR.get(from_net) or _CABLE_COLOR.get(cable, '#adb5bd')
        dashes = False
        if cable == 'cat5e':
            dashes = [8, 4]
        elif cable == 'cat5':
            dashes = [3, 3]
        edges.append({
            'id':    lk.id,
            'from':  lk.from_device_id,
            'to':    lk.to_device_id,
            'color': {'color': clr, 'highlight': clr, 'opacity': 0.9},
            'width': 3 if cable == 'fiber' else 2,
            'dashes': dashes,
            'title': f'케이블: {_CABLE_KO.get(cable, cable)}<br>망: {lk.network_type or from_net or "-"}',
            'label': _CABLE_KO.get(cable, ''),
            'font':  {'size': 9, 'align': 'middle', 'color': '#444'},
            # 필터용 메타
            'cable_type':   cable,
            'network_type': lk.network_type or from_net,
        })

    return {'nodes': nodes, 'edges': edges}


def generate_and_save_topology(school_id: int):
    """현재 장비/링크를 NetworkTopology 스냅샷으로 저장"""
    from apps.network.models import NetworkTopology
    from apps.schools.models import School

    school = School.objects.get(id=school_id)
    data   = build_topology_data(school_id)
    topo = NetworkTopology.objects.create(school=school, topology_data=data)
    try:
        write_topology_files_to_nas(school)
    except Exception:
        pass
    return topo


# ═══════════════════════════════════════════════════════
# NAS 파일 자동 생성 (토폴로지 CSV + SNMP 가이드 DOCX)
# ═══════════════════════════════════════════════════════

def _nas_topology_root():
    import os
    media = os.environ.get('NAS_MEDIA_ROOT', '/app/nas/media/npms')
    return os.path.join(media, '토폴로지')


def write_topology_csv(school, file_path: str):
    """장비 목록 CSV 생성"""
    import csv
    from apps.network.models import NetworkDevice
    TYPE_KO = {'switch':'스위치','poe_switch':'PoE스위치','ap':'무선AP','router':'라우터','firewall':'방화벽','server':'서버'}
    STATUS_KO = {'up':'정상','down':'장애','warning':'경고','unknown':'미확인'}
    devices = NetworkDevice.objects.filter(school=school).order_by('network_type', 'name')
    with open(file_path, 'w', encoding='utf-8-sig', newline='') as f:
        w = csv.writer(f)
        w.writerow(['장비명', '모델', '설치위치', '망구분', '장비유형', 'IP주소', '상태'])
        for d in devices:
            w.writerow([
                d.name, d.model, d.location, d.network_type,
                TYPE_KO.get(d.device_type, d.device_type),
                d.ip_address or '', STATUS_KO.get(d.status, d.status),
            ])


def write_snmp_guide_docx(school, file_path: str):
    """SNMP 설정 가이드 DOCX 생성 — 구성도 기반 상세 문서"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from django.utils import timezone
    from apps.network.models import NetworkDevice

    KF = '맑은 고딕'
    CF = 'Courier New'
    doc = Document()

    # ── 한글 폰트 헬퍼 ─────────────────────────
    def _setfont_style(st):
        st.font.name = KF
        rp = st.element.get_or_add_rPr()
        rf = rp.find(qn('w:rFonts'))
        if rf is None:
            rf = OxmlElement('w:rFonts'); rp.append(rf)
        for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
            rf.set(qn(a), KF)

    _setfont_style(doc.styles['Normal'])
    doc.styles['Normal'].font.size = Pt(10)
    for hn in ('Heading 1', 'Heading 2', 'Heading 3', 'Title'):
        try:
            _setfont_style(doc.styles[hn])
        except KeyError:
            pass

    def _apply(r):
        r.font.name = KF
        rp = r._element.get_or_add_rPr()
        rf = rp.find(qn('w:rFonts'))
        if rf is None:
            rf = OxmlElement('w:rFonts'); rp.append(rf)
        for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
            rf.set(qn(a), KF)

    # ── 색상 상수 ─────────────────────────────
    CLR_BLUE = RGBColor(0x1F, 0x5C, 0x99)
    CLR_BLUE2 = RGBColor(0x2E, 0x75, 0xB6)
    CLR_BODY = RGBColor(0x33, 0x33, 0x33)
    CLR_SUB = RGBColor(0x55, 0x55, 0x55)
    CLR_WARN = RGBColor(0x7A, 0x4F, 0x00)
    CLR_OK = RGBColor(0x15, 0x57, 0x24)
    CLR_CODE = RGBColor(0x1A, 0x1A, 0x1A)
    CLR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    FILL_HDR = '1F5C99'
    FILL_ROW_ODD = 'F5F9FF'
    FILL_ROW_EVEN = 'F8F8F8'

    def _set_cell_bg(cell, color_hex):
        """셀 배경색 설정"""
        tc = cell._element
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr'); tc.insert(0, tcPr)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color_hex)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    def add_p(t, style=None, color=None):
        p = doc.add_paragraph(t, style=style) if style else doc.add_paragraph(t)
        for r in p.runs:
            _apply(r)
            r.font.size = Pt(10)
            r.font.color.rgb = color or CLR_BODY
        return p

    def add_warn(t):
        """경고 문단 (갈색)"""
        return add_p(t, color=CLR_WARN)

    def add_ok(t):
        """안내 문단 (녹색)"""
        return add_p(t, color=CLR_OK)

    def add_h(t, lvl):
        h = doc.add_heading(t, level=lvl)
        for r in h.runs:
            _apply(r)
            if lvl == 1:
                r.font.size = Pt(16); r.font.color.rgb = CLR_BLUE
            elif lvl == 2:
                r.font.size = Pt(13); r.font.color.rgb = CLR_BLUE
            elif lvl == 3:
                r.font.size = Pt(11); r.font.color.rgb = CLR_BLUE2
        return h

    def add_h3(t):
        """Heading 3 with ▶ prefix"""
        return add_h(f'\u25B6 {t}', 3)

    def add_code(lines):
        """코드 블록 (Courier New 9pt, 들여쓰기, 진회색)"""
        for line in lines:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.8)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(line)
            run.font.name = CF
            run.font.size = Pt(9)
            run.font.color.rgb = CLR_CODE

    def add_table(headers, rows_data):
        """파란 헤더 + 교차 배경색 테이블"""
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = 'Table Grid'
        # 헤더: 파란 배경 + 흰 글씨
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            _set_cell_bg(hdr[i], FILL_HDR)
            for r in hdr[i].paragraphs[0].runs:
                _apply(r); r.font.bold = True; r.font.size = Pt(9)
                r.font.color.rgb = CLR_WHITE
        # 데이터 행: 교차 배경색
        for ri, row_vals in enumerate(rows_data):
            row = t.add_row().cells
            fill = FILL_ROW_ODD if ri % 2 == 0 else FILL_ROW_EVEN
            for i, v in enumerate(row_vals):
                row[i].text = str(v)
                _set_cell_bg(row[i], fill)
                for r in row[i].paragraphs[0].runs:
                    _apply(r); r.font.size = Pt(9)
                    r.font.color.rgb = CLR_BODY
        return t

    now_str = timezone.localtime(timezone.now()).strftime('%Y년 %m월 %d일')
    devices = list(NetworkDevice.objects.filter(school=school).order_by('network_type', 'name'))
    dev_cnt = len(devices)
    nets = sorted(set(d.network_type or '기타' for d in devices))
    fw_cnt = sum(1 for d in devices if d.device_type == 'firewall')
    sw_cnt = dev_cnt - fw_cnt

    # ══════════════════════════════════════════════
    # 표지
    # ══════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(f'{school.name} 네트워크')
    _apply(tr); tr.font.size = Pt(28); tr.font.bold = True; tr.font.color.rgb = CLR_BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run('NMS 개발을 위한 SNMP 설정 가이드')
    _apply(sr); sr.font.size = Pt(22); sr.font.bold = True; sr.font.color.rgb = CLR_BLUE2

    en = doc.add_paragraph()
    en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    er = en.add_run('Network Management System SNMP Integration Manual')
    _apply(er); er.font.size = Pt(12); er.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()
    info_rows = [
        ('대상 시스템', f'{school.name} 네트워크 구성도'),
        ('대상 장비 수', f'스위치 {sw_cnt}대 + 방화벽 {fw_cnt}대 (총 {dev_cnt}대)'),
        ('망 구성', ' / '.join(nets)),
        ('SNMP 버전', 'SNMPv2c (권장) / SNMPv3 (보안 강화)'),
        ('작성 목적', 'NMS 토폴로지 연동 및 실시간 모니터링 구현'),
        ('작성일', now_str),
    ]
    add_table(['항목', '내용'], info_rows)
    doc.add_page_break()

    # ══════════════════════════════════════════════
    # 1. SNMP 개요 및 NMS 연동 구조
    # ══════════════════════════════════════════════
    add_h('1. SNMP 개요 및 NMS 연동 구조', 1)

    add_h('1.1 SNMP란?', 2)
    add_p('SNMP(Simple Network Management Protocol)는 네트워크 장비를 원격으로 모니터링하고 '
          '관리하기 위한 표준 프로토콜입니다. NMS(Network Management System)는 SNMP를 통해 '
          '스위치, 방화벽, AP 등의 실시간 상태 정보를 수집합니다.')

    add_h3('SNMP 버전 비교')
    add_table(
        ['버전', '인증 방식', '보안 수준', '권장 용도'],
        [
            ['v1', 'Community String', '매우 낮음 (평문)', '미사용 권장'],
            ['v2c', 'Community String', '낮음 (평문)', '내부망 모니터링 (권장)'],
            ['v3', 'Username + Auth + Priv', '높음 (암호화)', '보안 강화 구간 (방화벽 등)'],
        ]
    )

    add_h('1.2 NMS 연동 동작 방식', 2)
    add_p('NMS가 SNMP로 장비를 모니터링하는 방식은 두 가지입니다:')
    add_p('Polling (주기적 수집): NMS가 일정 주기(예: 30초, 5분)마다 장비에 GET 요청을 보내 상태값을 수집',
          style='List Bullet')
    add_p('Trap (이벤트 수신): 장비에서 이상 발생 시 즉시 NMS로 알림 전송 (linkDown, coldStart 등)',
          style='List Bullet')
    doc.add_paragraph()
    add_warn('\u26A0 이 구성도의 모든 스위치/방화벽은 SNMP Agent 역할을 하며, '
             'NMS가 Manager 역할을 수행합니다.')

    # ══════════════════════════════════════════════
    # 2. 구성도 장비 목록 (SNMP 설정 대상)
    # ══════════════════════════════════════════════
    add_h('2. 구성도 장비 목록 (SNMP 설정 대상)', 1)
    add_p(f'아래 장비들이 {school.name} 네트워크 구성도에서 식별된 SNMP 모니터링 대상입니다. '
          '각 장비에 IP 주소를 부여하고 SNMP를 활성화해야 합니다.')

    TIER_MAP = {'firewall': '보안장비', 'router': '코어', 'switch': '', 'poe_switch': 'PoE',
                'ap': 'AP', 'server': '서버'}
    dev_rows = []
    for i, d in enumerate(devices, 1):
        tier = TIER_MAP.get(d.device_type, '')
        if not tier and d.name:
            # 계위 추정: 이름에 #M이면 1계위, #2~#5면 2계위 등
            if '#M' in d.name or '#BB' in d.name:
                tier = '1계위'
            else:
                tier = '스위치'
        dev_rows.append([
            str(i), d.name or '-', d.model or '-', d.location or '-',
            d.network_type or '-', tier, d.ip_address or '미등록',
        ])
    add_table(['No', '장비 ID', '모델명', '설치 위치', '망 구분', '계층', 'IP주소'], dev_rows)

    poe_devices = [d for d in devices if d.device_type == 'poe_switch']
    if poe_devices:
        doc.add_paragraph()
        poe_names = ', '.join(d.name for d in poe_devices[:5])
        if len(poe_devices) > 5:
            poe_names += f' 외 {len(poe_devices)-5}대'
        add_ok(f'\u2714 {poe_names}은(는) PoE 스위치입니다. PoE MIB(RFC 3621)를 추가로 활성화하면 '
               '무선AP 전력 상태도 NMS에서 모니터링할 수 있습니다.')

    # ══════════════════════════════════════════════
    # 3. 장비별 SNMP 설정 방법
    # ══════════════════════════════════════════════
    add_h('3. 장비별 SNMP 설정 방법', 1)

    add_h('3.1 Cisco/Allied Telesis 계열 스위치 (CBS220, GS724T, SG300 등)', 2)
    add_h3('SNMPv2c 설정 (CLI)')
    add_code([
        '! SNMP Community 설정 (읽기 전용)',
        'snmp-server community NMS_READ_ONLY ro',
        '',
        '! SNMP Trap 수신 NMS 서버 주소 설정',
        'snmp-server host 10.x.x.x version 2c NMS_READ_ONLY',
        '',
        '! Trap 이벤트 활성화',
        'snmp-server enable traps snmp linkdown linkup coldstart',
        'snmp-server enable traps port-security',
        '',
        '! 시스템 정보 등록 (NMS 식별용)',
        f'snmp-server location {school.name}',
        'snmp-server contact 네트워크관리자',
    ])

    add_h3('SNMPv3 설정 (보안 강화, 방화벽 권장)')
    add_code([
        '! SNMPv3 그룹 생성',
        'snmp-server group NMS_GROUP v3 priv',
        '',
        '! SNMPv3 사용자 생성 (SHA 인증 + AES 암호화)',
        'snmp-server user nmsuser NMS_GROUP v3',
        '  auth sha AuthPassword123',
        '  priv aes 128 PrivPassword456',
        '',
        '! NMS 서버 Trap 대상 설정',
        'snmp-server host 10.x.x.x version 3 priv nmsuser',
    ])

    fw_devices = [d for d in devices if d.device_type == 'firewall']
    if fw_devices:
        fw_models = ', '.join(set(d.model or '방화벽' for d in fw_devices))
        add_h(f'3.2 {fw_models} 방화벽', 2)
        add_p('방화벽 관리 웹 UI 또는 CLI에서 아래 항목을 설정합니다:')
        for s in [
            '관리 메뉴 > SNMP 설정 이동',
            'SNMP 에이전트 활성화 (Enable)',
            'Community String: NMS_READ_ONLY (읽기 전용)',
            '허용 NMS IP 등록: NMS 서버 IP 입력',
            'Trap 수신 주소: NMS 서버 IP, Port 162',
            'Trap 이벤트: CPU과부하, 세션초과, linkDown, 정책위반 등 활성화',
        ]:
            add_p(s, style='List Bullet')
        add_warn('\u26A0 방화벽은 SNMPv3 사용을 강력 권장합니다. '
                 'Community String 방식은 평문 전송으로 보안에 취약합니다.')

    if poe_devices:
        add_h('3.3 PoE 스위치 추가 설정', 2)
        add_code([
            '! PoE MIB 활성화 (IEEE 802.3af/at)',
            'snmp-server enable traps poe',
            '',
            '! PoE 포트 상태 Trap 설정',
            'interface GigabitEthernet 1/0/1',
            '  power inline auto',
            '  snmp trap poe-notification',
        ])

    # ══════════════════════════════════════════════
    # 4. NMS 개발용 주요 OID 목록
    # ══════════════════════════════════════════════
    add_h('4. NMS 개발용 주요 OID 목록', 1)
    add_p('아래 OID들을 NMS 소프트웨어에 등록하여 SNMP GET/GETNEXT/WALK로 수집합니다.')

    oid_data = [
        ['시스템 정보', 'sysDescr', '1.3.6.1.2.1.1.1.0', '장비 설명/모델명'],
        ['시스템 정보', 'sysName', '1.3.6.1.2.1.1.5.0', '장비 호스트명'],
        ['시스템 정보', 'sysUpTime', '1.3.6.1.2.1.1.3.0', '장비 업타임(timeticks)'],
        ['시스템 정보', 'sysContact', '1.3.6.1.2.1.1.4.0', '관리자 연락처'],
        ['시스템 정보', 'sysLocation', '1.3.6.1.2.1.1.6.0', '장비 설치 위치'],
        ['인터페이스', 'ifNumber', '1.3.6.1.2.1.2.1.0', '전체 포트 수'],
        ['인터페이스', 'ifDescr', '1.3.6.1.2.1.2.2.1.2', '포트 이름/설명'],
        ['인터페이스', 'ifOperStatus', '1.3.6.1.2.1.2.2.1.8', '포트 UP/DOWN 상태 (1=UP, 2=DOWN)'],
        ['인터페이스', 'ifAdminStatus', '1.3.6.1.2.1.2.2.1.7', '포트 관리 상태'],
        ['인터페이스', 'ifSpeed', '1.3.6.1.2.1.2.2.1.5', '포트 속도 (bps)'],
        ['트래픽', 'ifHCInOctets', '1.3.6.1.2.1.31.1.1.1.6', '수신 바이트 (64bit)'],
        ['트래픽', 'ifHCOutOctets', '1.3.6.1.2.1.31.1.1.1.10', '송신 바이트 (64bit)'],
        ['트래픽', 'ifInErrors', '1.3.6.1.2.1.2.2.1.14', '수신 에러 수'],
        ['트래픽', 'ifOutErrors', '1.3.6.1.2.1.2.2.1.20', '송신 에러 수'],
        ['CPU/메모리', 'cpmCPUTotal5minRev', '1.3.6.1.4.1.9.9.109.1.1.1.1.8', 'CPU 5분 평균 (Cisco)'],
        ['CPU/메모리', 'ciscoMemoryPoolFree', '1.3.6.1.4.1.9.9.48.1.1.1.6', '여유 메모리 (Cisco)'],
        ['PoE', 'pethMainPseOperStatus', '1.3.6.1.2.1.105.1.3.1.1.3', 'PoE 동작 상태'],
        ['PoE', 'pethMainPseConsumptionPower', '1.3.6.1.2.1.105.1.3.1.1.4', 'PoE 소비 전력(W)'],
        ['PoE', 'pethPsePortDetectionStatus', '1.3.6.1.2.1.105.1.1.1.6', 'PoE 포트별 PD 탐지'],
        ['MAC/ARP', 'dot1dTpFdbAddress', '1.3.6.1.2.1.17.4.3.1.1', 'MAC 주소 테이블'],
        ['VLAN', 'vtpVlanName', '1.3.6.1.4.1.9.9.46.1.3.1.1.4', 'VLAN 이름 (Cisco 계열)'],
    ]
    add_table(['분류', 'OID 이름', 'OID 번호', '설명'], oid_data)
    doc.add_paragraph()
    add_ok('\u2714 Gigabit 이상 고속 링크는 ifInOctets/ifOutOctets(32bit) 대신 '
           'ifHCInOctets/ifHCOutOctets(64bit)를 사용해야 카운터 오버플로우를 방지할 수 있습니다.')

    # ══════════════════════════════════════════════
    # 5. SNMP Trap 설정 및 수신 구조
    # ══════════════════════════════════════════════
    add_h('5. SNMP Trap 설정 및 수신 구조', 1)

    add_h('5.1 주요 Trap OID 목록', 2)
    add_table(
        ['Trap 이름', 'OID', '발생 조건', '심각도'],
        [
            ['linkDown', '1.3.6.1.6.3.1.1.5.3', '포트 DOWN 발생', '긴급'],
            ['linkUp', '1.3.6.1.6.3.1.1.5.4', '포트 UP 복구', '정보'],
            ['coldStart', '1.3.6.1.6.3.1.1.5.1', '장비 콜드 부팅', '경고'],
            ['warmStart', '1.3.6.1.6.3.1.1.5.2', '장비 웜 부팅(재시작)', '경고'],
            ['authenticationFailure', '1.3.6.1.6.3.1.1.5.5', 'SNMP 인증 실패', '경고'],
            ['pethPsePortOnOffNotification', '1.3.6.1.2.1.105.0.1', 'PoE 포트 상태 변경', '정보'],
        ]
    )

    add_h('5.2 NMS Trap 수신 포트 및 처리 흐름', 2)
    add_p('NMS 서버는 UDP 162번 포트로 Trap을 수신합니다. 처리 흐름:')
    for line in [
        '장비에서 이벤트 발생',
        '  -> SNMP Trap 패킷 전송 (UDP 162 -> NMS 서버 IP)',
        '  -> NMS Trap Listener (snmptrapd 또는 직접 구현)',
        '  -> OID 파싱 -> 이벤트 분류 (linkDown / coldStart / authFailure ...)',
        '  -> DB 저장 + 토폴로지 상태 업데이트 + 알람 발송',
    ]:
        add_p(line)

    # ══════════════════════════════════════════════
    # 6. NMS 개발 구현 가이드
    # ══════════════════════════════════════════════
    add_h('6. NMS 개발 구현 가이드', 1)

    add_h('6.1 SNMP 라이브러리 추천', 2)
    add_table(
        ['언어', '라이브러리', '설치 명령', '특징'],
        [
            ['Python', 'pysnmp', 'pip install pysnmp', 'SNMP v1/v2c/v3, Trap 수신 지원'],
            ['Python', 'easysnmp', 'pip install easysnmp', 'Net-SNMP 래퍼, 고성능'],
            ['Java', 'SNMP4J', 'Maven: org.snmp4j', '엔터프라이즈급, 비동기 지원'],
            ['Node.js', 'net-snmp', 'npm install net-snmp', '경량, v1/v2c/v3 지원'],
            ['Go', 'gosnmp', 'go get github.com/gosnmp', '고성능, 대규모 폴링 적합'],
        ]
    )

    add_h('6.2 Python pysnmp 예제 코드', 2)

    add_h3('장비 상태 폴링 (SNMP GET)')
    add_code([
        'from pysnmp.hlapi import *',
        '',
        'def get_snmp_value(ip, community, oid):',
        '    iterator = getCmd(',
        '        SnmpEngine(),',
        '        CommunityData(community, mpModel=1),  # v2c',
        '        UdpTransportTarget((ip, 161), timeout=2, retries=1),',
        '        ContextData(),',
        '        ObjectType(ObjectIdentity(oid))',
        '    )',
        '    errorIndication, errorStatus, errorIndex, varBinds = next(iterator)',
        '    if errorIndication or errorStatus:',
        '        return None',
        '    return str(varBinds[0][1])',
        '',
        '# 사용 예시',
        "sysname  = get_snmp_value('10.0.0.10', 'NMS_READ_ONLY', '1.3.6.1.2.1.1.5.0')",
        "uptime   = get_snmp_value('10.0.0.10', 'NMS_READ_ONLY', '1.3.6.1.2.1.1.3.0')",
    ])

    add_h3('포트 상태 일괄 수집 (SNMP WALK)')
    add_code([
        'from pysnmp.hlapi import *',
        '',
        'def walk_interface_status(ip, community):',
        '    results = {}',
        '    for (errorIndication, errorStatus, errorIndex, varBinds) in nextCmd(',
        '        SnmpEngine(),',
        '        CommunityData(community, mpModel=1),',
        '        UdpTransportTarget((ip, 161)),',
        '        ContextData(),',
        "        ObjectType(ObjectIdentity('1.3.6.1.2.1.2.2.1.8')),  # ifOperStatus",
        '        lexicographicMode=False',
        '    ):',
        '        if errorIndication: break',
        '        for varBind in varBinds:',
        '            oid, value = varBind',
        "            port_idx = str(oid).split('.')[-1]",
        "            status = 'UP' if int(value) == 1 else 'DOWN'",
        '            results[port_idx] = status',
        '    return results',
    ])

    add_h3('Trap 수신 서버 구현')
    add_code([
        'from pysnmp.carrier.asyncore.dispatch import AsyncoreDispatcher',
        'from pysnmp.carrier.asyncore.dgram import udp',
        'from pyasn1.codec.ber import decoder',
        'from pysnmp.proto import api',
        '',
        'def callback(transportDispatcher, transportDomain, transportAddress, wholeMsg):',
        '    while wholeMsg:',
        '        msgVer = int(api.decodeMessageVersion(wholeMsg))',
        '        pMod = api.protoModules[msgVer]',
        '        reqMsg, wholeMsg = decoder.decode(wholeMsg, asn1Spec=pMod.Message())',
        '        reqPDU = pMod.apiMessage.getPDU(reqMsg)',
        "        print(f'[TRAP] 발신: {transportAddress[0]}')",
        '        for oid, val in pMod.apiPDU.getVarBinds(reqPDU):',
        "            print(f'  OID: {oid}, Value: {val}')",
        '        # -> DB 저장 / 알람 처리 로직 추가',
        '    return wholeMsg',
        '',
        'transportDispatcher = AsyncoreDispatcher()',
        'transportDispatcher.registerRecvCbFun(callback)',
        'transportDispatcher.registerTransport(',
        '    udp.domainName,',
        "    udp.UdpSocketTransport().openServerMode(('0.0.0.0', 162))",
        ')',
        "print('Trap 수신 대기 중... (UDP 162)')",
        'transportDispatcher.jobStarted(1)',
        'transportDispatcher.runDispatcher()',
    ])

    # ══════════════════════════════════════════════
    # 7. NMS 토폴로지 구현 시 고려사항
    # ══════════════════════════════════════════════
    add_h('7. NMS 토폴로지 구현 시 고려사항', 1)

    add_h('7.1 구성도를 토폴로지에 반영하는 방법', 2)
    add_p(f'{school.name} 구성도는 계층(Tier) 구조입니다. NMS 토폴로지에 아래 방식으로 반영하세요:')
    add_table(
        ['계층', '장비', 'NMS 구현 방법'],
        [
            ['보안 (Top)', f'방화벽 {fw_cnt}대', '최상위 노드로 배치. WAN 트래픽/세션 수 모니터링'],
            ['1계위', '망별 핵심 스위치 (#M, #BB)', 'VLAN 별 트래픽 분리 표시'],
            ['2계위', '건물/층별 분배 스위치', 'PoE 전력/AP 연결 상태 표시'],
            ['3~4계위', '교실/교무실 접속 스위치', '포트 UP/DOWN 중점 모니터링'],
        ]
    )

    add_h('7.2 SNMP 데이터 -> 토폴로지 상태 매핑', 2)
    for s in [
        'ifOperStatus = 1 (UP) -> 토폴로지 링크 녹색 표시',
        'ifOperStatus = 2 (DOWN) -> 토폴로지 링크 빨간색 + 알람 발생',
        'linkDown Trap 수신 -> 즉시 해당 노드/링크 상태 빨간색으로 변경',
        '트래픽 사용률 (ifHCInOctets/ifSpeed x 100) -> 링크 색상 그라데이션으로 표현',
        'coldStart Trap -> 장비 재부팅 알림 (주황색 경보)',
    ]:
        add_p(s, style='List Bullet')

    add_h('7.3 Polling 주기 권장값', 2)
    add_table(
        ['수집 항목', '권장 주기', '비고'],
        [
            ['포트 UP/DOWN 상태 (ifOperStatus)', '30초', '빠른 장애 감지용'],
            ['트래픽 사용률 (ifHCInOctets)', '5분', '표준 모니터링 주기'],
            ['시스템 정보 (sysDescr, sysName)', '1시간', '장비 정보는 거의 변하지 않음'],
            ['CPU/메모리 사용률', '1분', '임계치(80%) 초과 시 알람'],
            ['PoE 전력 소비량', '5분', 'PoE 스위치 해당'],
        ]
    )

    # ══════════════════════════════════════════════
    # 8. 보안 권고사항
    # ══════════════════════════════════════════════
    add_h('8. 보안 권고사항', 1)
    for s in [
        "Community String은 기본값 'public', 'private' 절대 사용 금지 -> 고유한 문자열 사용",
        'SNMP ACL 적용: NMS 서버 IP만 SNMP 응답 허용 (access-list로 제한)',
        '방화벽에서 UDP 161(SNMP), UDP 162(Trap)을 NMS 서버 IP만 허용',
        '방화벽은 반드시 SNMPv3 (SHA+AES) 사용 권장',
        '학생망 장비는 교사망 NMS와 VLAN 분리 후 관리 포트로만 접근',
    ]:
        add_p(s, style='List Bullet')

    if poe_devices:
        doc.add_paragraph()
        add_warn('\u26A0 무선망 PoE 스위치의 SNMP 설정 시 관리 VLAN IP와 데이터 VLAN IP를 '
                 '혼동하지 않도록 주의하세요.')

    # ══════════════════════════════════════════════
    # 9. NMS 구축 체크리스트
    # ══════════════════════════════════════════════
    add_h('9. NMS 구축 체크리스트', 1)
    checklist = [
        '전체 장비에 관리용 IP 주소 부여 및 ping 연결 확인',
        '각 장비 CLI/웹UI에서 SNMP 에이전트 활성화',
        'Community String 설정 (기본값 변경 필수)',
        'SNMP Trap 수신 서버(NMS IP) 등록',
        '방화벽 SNMP ACL(허용IP) 설정',
        'NMS 서버에서 snmpget/snmpwalk 테스트',
        '주요 OID 수집 확인 (sysName, ifOperStatus 등)',
        'Trap 수신 테스트 (케이블 분리 -> linkDown 확인)',
        'NMS 토폴로지에 장비 배치 및 링크 설정',
        'PoE 장비 MIB 활성화 및 전력 모니터링 확인',
        '알람 룰 설정 (포트 DOWN, CPU 80% 초과 등)',
        '알람 임계치 설정 (CPU 80%, 포트 DOWN 즉시 알람)',
        'PoE 장비 전력 모니터링 확인',
    ]
    add_table(
        ['No', '작업 항목', '담당', '완료'],
        [[str(i), s, '', ''] for i, s in enumerate(checklist, 1)]
    )

    doc.add_paragraph()
    add_p(f'본 문서는 {school.name} NMS 개발 참고용으로 작성되었습니다.')

    doc.save(file_path)


def write_topology_files_to_nas(school):
    """학교의 토폴로지 CSV + SNMP 가이드 DOCX를 NAS에 자동 생성
    저장 위치:
      /app/nas/media/npms/토폴로지/토폴로지/토폴로지_{학교명}_{YYYYMMDD}.csv
      /app/nas/media/npms/토폴로지/SNMP설정가이드/SNMP설정가이드_{학교명}_{YYYYMMDD}.docx
    """
    import os
    import logging
    from datetime import datetime
    log = logging.getLogger(__name__)
    root = _nas_topology_root()
    topo_dir = os.path.join(root, '토폴로지')
    snmp_dir = os.path.join(root, 'SNMP설정가이드')
    os.makedirs(topo_dir, exist_ok=True)
    os.makedirs(snmp_dir, exist_ok=True)
    today = datetime.now().strftime('%Y%m%d')
    csv_path = os.path.join(topo_dir, f'토폴로지_{school.name}_{today}.csv')
    docx_path = os.path.join(snmp_dir, f'SNMP설정가이드_{school.name}_{today}.docx')
    try:
        write_topology_csv(school, csv_path)
    except Exception as e:
        log.warning(f'[{school.name}] CSV 생성 실패: {e}')
    try:
        write_snmp_guide_docx(school, docx_path)
    except Exception as e:
        log.warning(f'[{school.name}] DOCX 생성 실패: {e}')
    return {'csv': csv_path, 'docx': docx_path}
